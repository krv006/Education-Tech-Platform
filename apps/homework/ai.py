"""Gemini bilan uy vazifasini baholash — AI-home-checker (RJalol) porti.

Farqlar (maktab platformasi uchun moslashuv):
  - Rol universitet emas, MAKTAB o'qituvchisi; feedback matnlari O'ZBEK tilida.
  - Fan kursning `subject` matnidan avtomatik aniqlanadi (detect_profile).
  - JSON sxema asl loyihadagi bilan bir xil qoldirilgan:
      {overall_score, grade, questions[], summary} — hisobotlar mos bo'ladi.

Gemini kaliti settings.GEMINI_API_KEY dan olinadi (env: GEMINI_API_KEY).
PDF/rasm/audio to'g'ridan-to'g'ri multimodal yuboriladi; DOCX matni lokal
ajratiladi (python-docx). google-generativeai importi kech (lazy) — kalit
ishlatilmaydigan joylarda (testlar, boshqa app'lar) paket talab qilinmaydi.
"""
import json
import re
import time
from dataclasses import dataclass
from pathlib import Path

from django.conf import settings

INLINE_SIZE_LIMIT_BYTES = 15 * 1024 * 1024  # shundan katta fayl File API orqali

ALLOWED_EXTENSIONS = {
    '.pdf', '.png', '.jpg', '.jpeg', '.webp', '.docx',
    '.mp3', '.wav', '.m4a', '.ogg',
}
AUDIO_EXTENSIONS = {'.mp3', '.wav', '.m4a', '.ogg'}
MAX_FILE_SIZE_MB = 25

GENERATION_CONFIG = {
    'temperature': 0.1,  # past temperatura — barqaror, qat'iy baholash
    'top_p': 0.9,
    'max_output_tokens': 8192,
    'response_mime_type': 'application/json',
}

# 0-100 ball -> baho yorlig'i (asl loyihadagi GRADE_SCALE)
GRADE_SCALE = [
    (90, "A'lo"),
    (80, 'Juda yaxshi'),
    (70, 'Yaxshi'),
    (60, 'Qoniqarli'),
    (50, 'Yaxshilash kerak'),
    (0, "Jiddiy yaxshilash kerak"),
]


class HomeworkAIError(Exception):
    pass


class InvalidModelResponseError(HomeworkAIError):
    pass


# ---------------------------------------------------------------------------
# Fan profillari — asl prompts.py dan maktab darajasiga qisqartirilgan port
# ---------------------------------------------------------------------------
@dataclass
class SubjectProfile:
    key: str
    role: str
    objectives: list
    error_categories: list
    rubric: str


_COMMON_OBJECTIVES = [
    "Read the student's work carefully and understand their reasoning.",
    'Determine whether each final answer is correct.',
    'Check every intermediate step.',
    'Identify conceptual, calculation, and logic mistakes.',
    'Identify missing steps.',
]

SUBJECTS = {
    'math': SubjectProfile(
        key='math',
        role='an expert school Mathematics teacher with 20+ years of experience grading homework',
        objectives=_COMMON_OBJECTIVES + [
            'Identify arithmetic, algebra, sign, and notation mistakes.',
        ],
        error_categories=[
            'Conceptual Error', 'Calculation Error', 'Algebra Error', 'Sign Error',
            'Notation Error', 'Formula Error', 'Logic Error', 'Missing Step', 'Arithmetic Error',
        ],
        rubric=(
            '100 = perfect\n90-99 = minor notation issues\n'
            '70-89 = mostly correct but contains mathematical mistakes\n'
            '40-69 = partially correct\n1-39 = major conceptual misunderstanding\n0 = no valid attempt'
        ),
    ),
    'physics': SubjectProfile(
        key='physics',
        role='an expert school Physics teacher with 20+ years of experience grading problem sets',
        objectives=_COMMON_OBJECTIVES + [
            'Check units and dimensional analysis.',
            'Check the physical setup and assumptions.',
        ],
        error_categories=[
            'Conceptual Error', 'Calculation Error', 'Unit Error', 'Formula Error',
            'Sign Error', 'Missing Step', 'Incorrect Assumption', 'Logic Error',
        ],
        rubric=(
            '100 = perfect, correct units and reasoning\n90-99 = minor unit or rounding issues\n'
            '70-89 = mostly correct physics but calculation mistakes\n'
            '40-69 = partially correct setup\n1-39 = major conceptual misunderstanding\n0 = no valid attempt'
        ),
    ),
    'chemistry': SubjectProfile(
        key='chemistry',
        role='an expert school Chemistry teacher with 20+ years of experience grading homework',
        objectives=_COMMON_OBJECTIVES + [
            'Check balanced equations, formulas, and nomenclature.',
            'Check stoichiometry calculations and units.',
        ],
        error_categories=[
            'Conceptual Error', 'Calculation Error', 'Stoichiometry Error', 'Formula Error',
            'Nomenclature Error', 'Unit Error', 'Unbalanced Equation', 'Missing Step',
        ],
        rubric=(
            '100 = perfect, correctly balanced and reasoned\n90-99 = minor issues\n'
            '70-89 = mostly correct chemistry but calculation mistakes\n'
            '40-69 = partially correct\n1-39 = major misunderstanding\n0 = no valid attempt'
        ),
    ),
    'biology': SubjectProfile(
        key='biology',
        role='an expert school Biology teacher with 20+ years of experience grading homework',
        objectives=[
            "Read the student's answer carefully.",
            'Determine whether the answer is scientifically accurate.',
            'Check terminology and described processes.',
            'Identify factual and conceptual mistakes.',
            'Identify missing steps in described processes.',
        ],
        error_categories=[
            'Conceptual Error', 'Factual Error', 'Terminology Error', 'Missing Step',
            'Incorrect Assumption', 'Incomplete Explanation',
        ],
        rubric=(
            '100 = accurate and complete\n90-99 = minor terminology issues\n'
            '70-89 = mostly correct but factual mistakes\n40-69 = partially correct\n'
            '1-39 = major misunderstanding\n0 = no valid attempt'
        ),
    ),
    'computer_science': SubjectProfile(
        key='computer_science',
        role='an expert school Informatics/Programming teacher with 20+ years of experience grading code',
        objectives=[
            "Read the student's code or answer carefully.",
            'Understand the algorithmic approach.',
            'Determine whether the solution works on typical and edge-case inputs.',
            'Identify logic errors and bugs.',
            'Identify missing edge-case handling and style issues.',
        ],
        error_categories=[
            'Logic Error', 'Algorithm Error', 'Syntax Error', 'Edge Case Missing',
            'Style Issue', 'Missing Step', 'Inefficient Solution',
        ],
        rubric=(
            '100 = correct and well-structured\n90-99 = minor style issues\n'
            '70-89 = mostly correct but logic bugs\n40-69 = partially correct approach\n'
            '1-39 = major misunderstanding\n0 = no valid attempt'
        ),
    ),
    'essay': SubjectProfile(
        key='essay',
        role='an expert Language Arts and Literature teacher with 20+ years of experience grading essays',
        objectives=[
            "Read the student's essay carefully.",
            'Determine whether the thesis/main idea is clear and supported.',
            'Check structure, organization, and transitions.',
            'Identify grammar, punctuation, and style issues.',
            'Identify unsupported or inconsistent claims.',
        ],
        error_categories=[
            'Weak Thesis', 'Unsupported Claim', 'Structure Issue', 'Grammar Error',
            'Punctuation Error', 'Logic Error', 'Clarity Issue', 'Style Issue',
        ],
        rubric=(
            '100 = excellent argument and mechanics\n90-99 = minor style issues\n'
            '70-89 = solid but structural weaknesses\n40-69 = underdeveloped\n'
            '1-39 = unclear or unsupported\n0 = no valid attempt'
        ),
    ),
    'history': SubjectProfile(
        key='history',
        role='an expert History teacher with 20+ years of experience grading written answers',
        objectives=[
            "Read the student's answer carefully.",
            'Determine whether facts, dates, and events are accurate.',
            'Check the argument, causation, and context.',
            'Identify factual errors and missing evidence.',
        ],
        error_categories=[
            'Factual Error', 'Chronology Error', 'Weak Argument', 'Missing Context',
            'Logic Error', 'Missing Evidence', 'Clarity Issue',
        ],
        rubric=(
            '100 = accurate and well-argued\n90-99 = very minor gaps\n'
            '70-89 = mostly accurate with weaknesses\n40-69 = partially accurate\n'
            '1-39 = largely inaccurate\n0 = no valid attempt'
        ),
    ),
    'general': SubjectProfile(
        key='general',
        role='an expert school teacher with 20+ years of experience grading homework across many subjects',
        objectives=_COMMON_OBJECTIVES,
        error_categories=[
            'Conceptual Error', 'Factual Error', 'Logic Error', 'Missing Step',
            'Incorrect Assumption', 'Clarity Issue',
        ],
        rubric=(
            '100 = perfect\n90-99 = minor issues only\n70-89 = mostly correct but notable mistakes\n'
            '40-69 = partially correct\n1-39 = major misunderstanding\n0 = no valid attempt'
        ),
    ),
}

# Til fanlari — (til, ko'nikma) juftligidan profil quriladi (asl SKILL_TEMPLATES porti)
LANGUAGES = {'english': 'English', 'russian': 'Russian', 'turkish': 'Turkish'}
SKILLS = {
    'writing': {
        'note': 'The submission is a piece of student writing (essay, letter, or short answers) as PDF, image, or Word document.',
        'objectives': [
            'Determine whether the task was fully addressed.',
            'Check grammar, verb tenses, and sentence structure.',
            'Check vocabulary range and spelling.',
            'Evaluate coherence and organization.',
        ],
        'error_categories': [
            'Grammar Error', 'Verb Tense Error', 'Vocabulary Error', 'Spelling Error',
            'Punctuation Error', 'Coherence Issue', 'Task Achievement Issue', 'Word Order Error',
        ],
        'rubric': (
            '100 = fully addresses the task, excellent language\n90-99 = very minor slips\n'
            '70-89 = noticeable grammar/vocabulary errors\n40-69 = frequent errors\n'
            '1-39 = task barely addressed\n0 = no valid attempt'
        ),
    },
    'reading': {
        'note': "The submission is a completed reading-comprehension worksheet (passage + the student's answers).",
        'objectives': [
            'Determine whether each answer is textually correct.',
            'Check main-idea and detail answers.',
            'Check inference answers against the text.',
        ],
        'error_categories': [
            'Comprehension Error', 'Incorrect Answer', 'Missing Detail',
            'Inference Error', 'Vocabulary Misunderstanding',
        ],
        'rubric': (
            '100 = all answers correct\n90-99 = one minor slip\n70-89 = some comprehension errors\n'
            '40-69 = several incorrect\n1-39 = most incorrect\n0 = no valid attempt'
        ),
    },
    'listening': {
        'note': "The submission is the student's completed listening-comprehension answer sheet.",
        'objectives': [
            'Determine whether each answer matches what was asked.',
            'Check spelling of dictated words.',
            'Identify missing or blank answers.',
        ],
        'error_categories': [
            'Comprehension Error', 'Incorrect Answer', 'Spelling Error', 'Missed Detail', 'Blank Answer',
        ],
        'rubric': (
            '100 = all correct and complete\n90-99 = very minor slip\n70-89 = a few details missed\n'
            '40-69 = several incorrect\n1-39 = little comprehension\n0 = no valid attempt'
        ),
    },
    'speaking': {
        'note': (
            'The submission is an AUDIO RECORDING of the student speaking. '
            'Listen to it directly, briefly transcribe what was said, then grade.'
        ),
        'objectives': [
            'Determine whether the task was addressed.',
            'Evaluate pronunciation and intelligibility.',
            'Evaluate fluency (hesitations, pauses).',
            'Check grammar and vocabulary used in speech.',
        ],
        'error_categories': [
            'Pronunciation Error', 'Fluency Issue', 'Grammar Error', 'Vocabulary Error',
            'Task Achievement Issue', 'Hesitation/Filler Overuse',
        ],
        'rubric': (
            '100 = excellent pronunciation, fluency, grammar; zero errors\n'
            '90-99 = at most one trivial slip\n70-89 = clear errors present (2+ mistakes listed => max this band)\n'
            '40-69 = frequent errors or heavy hesitation\n1-39 = very hard to understand\n'
            '0 = no audio content\n'
            'IMPORTANT: confident delivery alone never justifies a high score — count errors first.'
        ),
    },
}

# course.subject matnidan fan kalitini aniqlash (board'dagi regex uslubi)
_SUBJECT_PATTERNS = [
    (r'matem|algebra|geometr', ('math', '', '')),
    (r'fizik', ('physics', '', '')),
    (r'kimyo|ximiy', ('chemistry', '', '')),
    (r'biolog', ('biology', '', '')),
    (r'informat|dastur|program|kompyut', ('computer_science', '', '')),
    (r'tarix|histor', ('history', '', '')),
    (r'adabiyot|ona tili|essay|insho', ('essay', '', '')),
    (r'ingliz|english', ('general', '', 'english')),
    (r'rus tili|russian', ('general', '', 'russian')),
    (r'turk', ('general', '', 'turkish')),
]


def detect_profile(subject_text: str) -> tuple:
    """course.subject -> (subject_key, custom_name, language_key)."""
    text = (subject_text or '').lower()
    for pattern, found in _SUBJECT_PATTERNS:
        if re.search(pattern, text):
            return found
    return ('general', subject_text or '', '')


# ---------------------------------------------------------------------------
# Prompt qurish — asl build_system_prompt'ning ixcham porti
# ---------------------------------------------------------------------------
_OUTPUT_FORMAT = """# OUTPUT FORMAT

Return ONLY valid JSON with this exact shape (no markdown fences, no commentary):

{
  "overall_score": 84,
  "grade": "Juda yaxshi",
  "questions": [
    {
      "question_number": 1,
      "question": "...",
      "student_answer": "...",
      "expected_solution": "...",
      "analysis": "...",
      "mistakes": ["..."],
      "error_categories": ["..."],
      "correct_answer": "...",
      "suggestions": ["..."],
      "difficulty": "Easy | Medium | Hard",
      "score": 85
    }
  ],
  "summary": {
    "strengths": ["..."],
    "weaknesses": ["..."],
    "topics_to_review": ["..."],
    "recommendations": ["..."]
  }
}

GRADING SCALE (for the "grade" field — use these exact Uzbek labels):
90-100 -> A'lo
80-89  -> Juda yaxshi
70-79  -> Yaxshi
60-69  -> Qoniqarli
50-59  -> Yaxshilash kerak
Below 50 -> Jiddiy yaxshilash kerak

LANGUAGE: write ALL human-readable text values (question, student_answer,
expected_solution, analysis, mistakes, correct_answer, suggestions, and every
summary item) in simple UZBEK (latin script) that a school student easily
understands. Keep JSON keys, "difficulty" values, and "error_categories" in
English exactly as listed.

The response must always be valid JSON. No markdown. Nothing outside JSON.
"""

_SCORING_DISCIPLINE = """# SCORING DISCIPLINE (STRICT)

1. First list every mistake in "mistakes", only then decide the score.
2. If "mistakes" is non-empty the score MUST reflect it — any real error means below 90.
3. Confident or long answers are not automatically good — grade only correctness.
4. Never round scores upward; if unsure between two bands choose the LOWER one.
5. A strong final answer does not excuse a flawed process, and vice versa.
"""


def build_system_prompt(subject_key: str, custom_name: str = '',
                        language_key: str = '', skill_key: str = '') -> str:
    if language_key and skill_key and skill_key in SKILLS:
        language = LANGUAGES.get(language_key, language_key.title())
        skill = SKILLS[skill_key]
        role = (
            f'an expert {language} language teacher and examiner with 20+ years '
            f'of experience assessing school students'
        )
        submission_note = skill['note']
        objectives = skill['objectives']
        error_categories = skill['error_categories']
        rubric = skill['rubric']
    else:
        profile = SUBJECTS.get(subject_key, SUBJECTS['general'])
        role = profile.role
        if profile.key == 'general' and custom_name.strip():
            role = (
                f'an expert {custom_name.strip()} teacher with 20+ years of '
                f'experience grading school homework'
            )
        submission_note = 'PDF, image (photo of the notebook page), or Word document; typed or handwritten.'
        objectives = profile.objectives
        error_categories = profile.error_categories
        rubric = profile.rubric

    objectives_block = '\n'.join(f'{i + 1}. {o}' for i, o in enumerate(objectives))
    return f"""You are {role}.

You grade SCHOOL-LEVEL homework. Be a strict, rigorous, fair examiner whose
scores can be trusted — kindness in tone, strictness in score.

# CONTEXT
The uploaded submission: {submission_note}
The file may contain multiple questions/tasks — detect every one automatically.

# OBJECTIVES (for every detected question/task)
{objectives_block}

# GRADING POLICY
Each question gets a correctness score 0-100 where:
{rubric}

{_SCORING_DISCIPLINE}

# FEEDBACK STYLE
Feedback is educational: explain what is wrong, why, and how to fix it.
Never insult the student. Never write just "Wrong."

# ERROR CATEGORIES
Use one or more of: {', '.join(error_categories)}

{_OUTPUT_FORMAT}"""


def build_user_prompt(extra_instructions: str = '') -> str:
    base = (
        'Analyze the attached homework submission according to your system '
        'instructions. Detect every question automatically, grade each one, '
        'and return ONLY the JSON object described in your instructions.'
    )
    if extra_instructions.strip():
        base += f"\n\nAdditional context from the teacher: {extra_instructions.strip()}"
    return base


# ---------------------------------------------------------------------------
# JSON tozalash/tekshirish — asl utils.py porti
# ---------------------------------------------------------------------------
def _strip_markdown_fences(text: str) -> str:
    cleaned = (text or '').strip()
    cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
    cleaned = re.sub(r'\s*```$', '', cleaned)
    return cleaned.strip()


def parse_and_validate_json(raw_text: str) -> dict:
    cleaned = _strip_markdown_fences(raw_text)
    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise InvalidModelResponseError(f'Model javobi JSON emas: {exc}') from exc

    missing = {'overall_score', 'grade', 'questions', 'summary'} - set(data)
    if missing:
        raise InvalidModelResponseError(f'JSONda majburiy kalitlar yetishmaydi: {sorted(missing)}')
    if not isinstance(data['questions'], list) or not data['questions']:
        raise InvalidModelResponseError("JSONda birorta ham savol yo'q ('questions' bo'sh).")

    required_q = {
        'question_number', 'question', 'student_answer', 'expected_solution',
        'analysis', 'mistakes', 'error_categories', 'correct_answer',
        'suggestions', 'difficulty', 'score',
    }
    for i, q in enumerate(data['questions']):
        missing_q = required_q - set(q)
        if missing_q:
            raise InvalidModelResponseError(
                f"{i + 1}-savolda kalitlar yetishmaydi: {sorted(missing_q)}"
            )
    return data


def grade_label(score) -> str:
    try:
        value = float(score)
    except (TypeError, ValueError):
        return ''
    for low, label in GRADE_SCALE:
        if value >= low:
            return label
    return GRADE_SCALE[-1][1]


_MIME_TYPES = {
    '.pdf': 'application/pdf',
    '.png': 'image/png',
    '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.webp': 'image/webp',
    '.mp3': 'audio/mpeg',
    '.wav': 'audio/wav',
    '.m4a': 'audio/mp4',
    '.ogg': 'audio/ogg',
}


def _extract_docx_text(path: Path) -> str:
    import docx  # python-docx — lazy, faqat DOCX kelganda kerak

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


# ---------------------------------------------------------------------------
# Asosiy kirish nuqtasi — asl HomeworkChecker.analyze porti
# ---------------------------------------------------------------------------
def grade_file(
    file_path,
    *,
    subject_text: str = '',
    skill_key: str = '',
    extra_instructions: str = '',
    max_retries: int = 2,
) -> dict:
    """Bitta topshiriq faylini Gemini bilan baholaydi, tekshirilgan JSON qaytaradi."""
    api_key = getattr(settings, 'GEMINI_API_KEY', '')
    if not api_key:
        raise HomeworkAIError(
            'GEMINI_API_KEY sozlanmagan — serverda env o\'zgaruvchisini bering.'
        )

    import google.generativeai as genai  # lazy — paket faqat shu yerda kerak

    subject_key, custom_name, language_key = detect_profile(subject_text)
    system_prompt = build_system_prompt(
        subject_key,
        custom_name=custom_name,
        language_key=language_key,
        skill_key=skill_key if language_key else '',
    )

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
        model_name=getattr(settings, 'GEMINI_MODEL', 'gemini-3.5-flash'),
        system_instruction=system_prompt,
    )

    path = Path(file_path)
    user_prompt = build_user_prompt(extra_instructions)
    ext = path.suffix.lower()
    if ext == '.docx':
        content_parts = [
            user_prompt,
            'The homework was submitted as a Word document. Extracted text:\n\n'
            + _extract_docx_text(path),
        ]
    else:
        mime_type = _MIME_TYPES.get(ext)
        try:
            if path.stat().st_size <= INLINE_SIZE_LIMIT_BYTES:
                content_parts = [user_prompt, {'mime_type': mime_type, 'data': path.read_bytes()}]
            else:
                content_parts = [user_prompt, genai.upload_file(path=str(path), mime_type=mime_type)]
        except Exception as exc:
            raise HomeworkAIError(f"'{path.name}' faylini Gemini'ga tayyorlab bo'lmadi: {exc}") from exc

    last_error = None
    for attempt in range(max_retries + 1):
        try:
            response = model.generate_content(content_parts, generation_config=GENERATION_CONFIG)
            return parse_and_validate_json(response.text)
        except InvalidModelResponseError as exc:
            last_error = exc
            content_parts = content_parts + [
                'Your previous response was not valid JSON matching the required '
                'schema. Return ONLY the corrected JSON object, nothing else.'
            ]
            time.sleep(1)
        except Exception as exc:  # tarmoq / API xatolari
            last_error = exc
            time.sleep(1.5 * (attempt + 1))

    raise HomeworkAIError(
        f"{max_retries + 1} urinishdan keyin ham yaroqli natija olinmadi: {last_error}"
    )
